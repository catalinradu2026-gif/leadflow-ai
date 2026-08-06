import re, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x0f,0x17,0x2a); BLUE=RGBColor(0x25,0x63,0xeb); GREY=RGBColor(0x64,0x74,0x8b)
KEEP={0x2013,0x2014,0x2018,0x2019,0x201C,0x201D,0x201E,0x2026}
def clean(s): return ''.join(c for c in s if ord(c)<0x2190 or ord(c) in KEEP)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def add_runs(p, text):
    text=clean(text)
    # split by bold **...**, then handle `code`
    for i,seg in enumerate(re.split(r'\*\*(.+?)\*\*', text)):
        if seg=='' : continue
        bold = (i%2==1)
        # links [t](u) -> t (bold)
        seg=re.sub(r'\[(.+?)\]\((.+?)\)', r'\1', seg)
        # code `x` -> keep text, mono
        parts=re.split(r'`(.+?)`', seg)
        for j,pt in enumerate(parts):
            if pt=='':continue
            r=p.add_run(pt)
            if bold: r.bold=True
            if j%2==1: r.font.name='Consolas'; r.font.size=Pt(9)

def build(mdfile, outfile, footer_text):
    lines=open(mdfile,encoding='utf-8').read().split('\n')
    doc=Document()
    st=doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
    for s in doc.sections:
        s.top_margin=Cm(1.8); s.bottom_margin=Cm(1.8); s.left_margin=Cm(2); s.right_margin=Cm(2)
        fp=s.footer.paragraphs[0]; fp.text=footer_text; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        for r in fp.runs: r.font.size=Pt(7.5); r.font.color.rgb=GREY
    i=0
    while i<len(lines):
        ln=lines[i].rstrip()
        if not ln.strip(): i+=1; continue
        if ln.startswith('# '):
            h=doc.add_heading('',level=0); r=h.add_run(clean(ln[2:])); r.font.color.rgb=NAVY; r.font.size=Pt(20)
        elif ln.startswith('## '):
            h=doc.add_heading('',level=1); r=h.add_run(clean(ln[3:])); r.font.color.rgb=BLUE; r.font.size=Pt(14)
        elif ln.startswith('### '):
            h=doc.add_heading('',level=2); r=h.add_run(clean(ln[4:])); r.font.color.rgb=NAVY; r.font.size=Pt(12)
        elif ln.startswith('> '):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
            r=p.add_run(clean(ln[2:])); r.italic=True; r.font.color.rgb=GREY; r.font.size=Pt(9.5)
        elif ln.strip().startswith('!['):
            m=re.match(r'!\[(.*?)\]\((.+?)\)', ln.strip())
            if m:
                cap, path = m.group(1), m.group(2)
                imgpath = path if os.path.isabs(path) else os.path.normpath(os.path.join(os.path.dirname(mdfile), path))
                p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                run=p.add_run(); pic=run.add_picture(imgpath)
                maxW, maxH = Cm(15.5), Cm(20)
                ratio=min(maxW/pic.width, maxH/pic.height)
                pic.width=Emu(int(pic.width*ratio)); pic.height=Emu(int(pic.height*ratio))
                if cap:
                    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=cp.add_run(clean(cap)); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GREY
            i+=1; continue
        elif ln.strip()=='---':
            doc.add_paragraph('_'*1).runs  # spacer
        elif ln.lstrip().startswith(('- ','* ')):
            while i<len(lines) and lines[i].lstrip().startswith(('- ','* ')):
                p=doc.add_paragraph(style='List Bullet'); add_runs(p, lines[i].lstrip()[2:]); i+=1
            continue
        elif re.match(r'^\d+\. ', ln.strip()):
            while i<len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
                p=doc.add_paragraph(style='List Number'); add_runs(p, re.sub(r'^\d+\. ','',lines[i].strip())); i+=1
            continue
        elif ln.strip().startswith('|'):
            rows=[]
            while i<len(lines) and lines[i].strip().startswith('|'):
                rows.append([c.strip() for c in lines[i].strip().strip('|').split('|')]); i+=1
            rows=[r for r in rows if not all(set(c)<=set('-: ') for c in r)]
            if rows:
                ncol=max(len(r) for r in rows)
                t=doc.add_table(rows=len(rows),cols=ncol); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
                for ri,row in enumerate(rows):
                    row=row+['']*(ncol-len(row))
                    for ci,cell in enumerate(row):
                        c=t.cell(ri,ci); c.paragraphs[0].text=''
                        add_runs(c.paragraphs[0], cell)
                        for run in c.paragraphs[0].runs: run.font.size=Pt(9)
                        if ri==0:
                            for run in c.paragraphs[0].runs: run.bold=True; run.font.color.rgb=RGBColor(0xff,0xff,0xff)
                            shade(c,'0f172a')
                doc.add_paragraph()
            continue
        else:
            p=doc.add_paragraph(); add_runs(p, ln)
        i+=1
    doc.save(outfile); print('OK ->', outfile.split('/')[-1])
